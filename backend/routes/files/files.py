import io
import re
import uuid

from fastapi import APIRouter, UploadFile, File
from docx import Document

from db.crud import LegendDB, MunicipalityDB
from db.db import client
from routes.legends.models import types_list

router = APIRouter(prefix="/files", tags=["Files"])


@router.post("/upload", name="Сохранить легенды файлом")
async def upload_file(file: bytes = File(...)):
    """ Сохранить легенды файлом """
    document = Document(io.BytesIO(file))
    for table in document.tables:
        for row in table.rows:
            if re.match("[0-9.]+$", row.cells[0].paragraphs[0].text):
                municipality = row.cells[1].paragraphs[0].text.strip().capitalize()
                result = await MunicipalityDB.get(name=municipality)
                if not result:
                    # Сохраняем новый муниципалитет, если его нет в БД
                    await MunicipalityDB.create(name=municipality, lat=58, long=32)
                    result = await MunicipalityDB.get(name=municipality)
                count = -1
                legend_cells = row.cells[2:11]
                info_cells = row.cells[11:]
                for cell in legend_cells:
                    count += 1
                    for para in cell.paragraphs:
                        if para.text and not re.match("[-]+$", para.text.strip()):
                            # Сохраняем легенду
                            await LegendDB.create(municipality_id=result.get("id"),
                                                  documents=''.join(p.text for p in info_cells[0].paragraphs),
                                                  description=para.text.strip().capitalize(), type=types_list[count],
                                                  informant=''.join(p.text for p in info_cells[1].paragraphs),
                                                  lat=result.get("lat", 58), long=result.get("long", 32))

    return


@router.post("/upload-audio", response_model=str, name="Загрузка файла для аудиогида")
def set_file(file: UploadFile = File(...)):
    """ Загрузка файла для аудиогида """
    file_id = uuid.uuid4().hex
    try:
        client.fput_object("audioguides", file_id, file.file.fileno())
    except:
        pass
    return file_id


@router.get("/get-audio/{file_id}", name="Получить файла для аудиогида")
def get_file(file_id: str):
    """ Получить файла для аудиогида """
    data = client.get_object("audioguides", file_id)
    return data
