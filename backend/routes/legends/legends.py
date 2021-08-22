import io
import re
import uuid

from fastapi import APIRouter, UploadFile, File
from docx import Document
from db.crud import LegendDB, MunicipalityDB
from db.db import client
from routes.legends.models import LegendResponseModel, LegendSaveRequestModel, types_list

router = APIRouter(prefix="/legends", tags=["Legends"])


@router.get("/get/{item_id}", response_model=LegendResponseModel, name="Получить легенду по ID")
async def get_legend(item_id: int):
    """ Получить легенду по ID """
    legend = await LegendDB.get(id=item_id)
    municipality = await MunicipalityDB.get(id=legend.get("municipality_id"))
    response = legend.copy()
    response["municipality"] = municipality.get("name")
    return response


@router.post("/save", name="Сохранить легенду")
async def set_legend(body: LegendSaveRequestModel):
    """ Сохранить легенду """
    response = await LegendDB.create(**body.dict())
    return response


@router.get("/get-all", response_model=list[LegendResponseModel], name="Получить список городов с легендами")
async def get_legends():
    """ Получить список легенд """
    legends = await LegendDB.search()
    municipalities = await MunicipalityDB.search()
    response = list()
    if legends:
        for legend in legends:
            for municipality in municipalities:
                if legend.get("municipality_id") == municipality.get("id"):
                    tmp = legend.copy()
                    tmp["municipality"] = municipality.get("name")
                    response.append(tmp)
        return response


@router.post("/upload-file", name="Сохранить легенды файлом")
async def set_file(file: bytes = File(...)):
    """ Сохранить легенды файлом """
    document = Document(io.BytesIO(file))
    for table in document.tables:
        for row in table.rows:
            if re.match("[0-9.]+$", row.cells[0].paragraphs[0].text):
                municipality = row.cells[1].paragraphs[0].text.strip().capitalize()
                result = await MunicipalityDB.get(name=municipality)
                if not result:
                    # Сохраняем новый муниципалитет, если его нет в БД
                    await MunicipalityDB.create(name=municipality, lat=58, long=32)
                    result = await MunicipalityDB.get(name=municipality)
                count = -1
                legend_cells = row.cells[2:11]
                info_cells = row.cells[11:]
                for cell in legend_cells:
                    count += 1
                    for para in cell.paragraphs:
                        if para.text and not re.match("[-]+$", para.text.strip()):
                            # Сохраняем легенду
                            await LegendDB.create(municipality_id=result.get("id"),
                                                  documents=''.join(p.text for p in info_cells[0].paragraphs),
                                                  description=para.text.strip().capitalize(), type=types_list[count],
                                                  informant=''.join(p.text for p in info_cells[1].paragraphs),
                                                  lat=result.get("lat", 58), long=result.get("long", 32))

    return


@router.post("/upload-audio-file/{legend_id}", name="Загрузка файла для аудиогида")
async def set_file(legend_id: int, file: UploadFile = File(...)):
    """ Загрузка файла для аудиогида """
    file_id = uuid.uuid4()
    await LegendDB.update(id=legend_id, audio_guide_id=file_id)
    client.fput_object("audioguides", file_id, file.file)
    return


@router.get("/get-audio-guide/{file_id}", name="Получить файла для аудиогида")
def get_file(file_id: int):
    """ Получить файла для аудиогида """
    data = client.get_object("audioguides", file_id)
    return data
